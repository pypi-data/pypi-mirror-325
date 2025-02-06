import json
import io
import concurrent.futures
import os
import requests
import subprocess
import pandas as pd
from datetime import timedelta, date
import numpy as np
import plotly.graph_objects as go
from tqdm import tqdm
from .model import KalmanFilter
from docx import Document
from sklearn.metrics import mean_squared_error
import scipy.optimize

home = os.path.expanduser("~")
filename = ".ligrc"
path = os.path.join(home, filename)
try:
    with open(path, "r") as file:
        content = file.read()
except FileNotFoundError as e:
    print(e)
    print("Please make sure that auth_key==<auth_key> is stored in above file.")
    print("Required for BAP login!")
    exit()

auth_key = content.strip().split("==")[1]

# Login API endpoint
login_api = "http://10.10.4.40:5000/webapi/auth.cgi"
login_params = {
    "api": "SYNO.API.Auth",
    "version": "3",
    "method": "login",
    "account": "harisankar",
    "passwd": "8P1GTw#f",
    "session": "FileStation",
    "format": "sid",
}


def get_imei(IMEI):
    """
    Function to get_imei from dashboard
    ```
    import correction_monitor as cm
    eg: imei=get_imei("MD0AIOALAA00638")
    ```
    """
    headers = headers = {
        "Authorization": f"{auth_key}",
        "Content-Type": "application/json",
    }
    response = requests.get(
        f"https://api-stage.lime.ai/lime/liveBatteryDashboard/{IMEI}",
        headers=headers,
    )
    try:
        return json.loads(response.content)["result"][0]["imei"]
    except Exception as e:
        raise ValueError("get_imei error: ", e)


def get_dates(start_date1, end_date1):
    """
    Returns a list of dates between the start and end dates, inclusive.

    Parameters:
    - start_date1: starting date
    - end_date1: ending date

    Returns:
    - List of date strings in the format 'YYYY-MM-DD'
    """
    start_day = int(start_date1.split("-")[2].strip())
    start_month = int(start_date1.split("-")[1].strip())
    start_year = int(start_date1.split("-")[0].strip())
    end_day = int(end_date1.split("-")[2].strip())
    end_month = int(end_date1.split("-")[1].strip())
    end_year = int(end_date1.split("-")[0].strip())
    # Convert start and end dates to datetime objects
    start_date = date((start_year), (start_month), (start_day))
    end_date = date((end_year), (end_month), (end_day))

    # List to hold date strings
    date_list = []

    # Iterate through each day in the range
    current_date = start_date
    while current_date <= end_date:
        # Format date as 'YYYY-MM-DD'
        date_str = current_date.strftime("%Y-%m-%d")
        date_list.append(date_str)

        # Move to the next day
        current_date += timedelta(days=1)

    return date_list


def get_extdata(IMEI, start_time, end_time, filter_data=False):
    """
    Function to import data from dashboard
    ```
    import lime_internal_getter as ig
    odf = ig.get_extdata(*("MD0AIOALAA00638", "2024-06-27", "2024-06-28"))
    ```
    """
    start_date = start_time.split(" ")[0]
    start_date = pd.to_datetime(start_date).strftime("%Y-%m-%d")
    end_date = end_time.split(" ")[0]
    end_date = pd.to_datetime(end_date).strftime("%Y-%m-%d")
    headers = headers = {
        "Authorization": f"{auth_key}",
        "Content-Type": "application/json",
    }
    payload = json.dumps({"imei": IMEI, "startDate": start_date, "endDate": end_date})
    response = requests.post(
        "https://api-stage.lime.ai/lime/iotData", headers=headers, data=payload
    )
    try:
        df = pd.DataFrame(json.loads(response.content)["result"])
        if filter_data is True:
            df["timeStamp"] = pd.to_datetime(df["timeStamp"])
            df = (
                df[(df["timeStamp"] >= start_time) & (df["timeStamp"] <= end_time)]
            ).reset_index(drop=True)
        return df
    except Exception:
        raise ValueError("Cannot Form Dataframe")


def get_pimdata(
    IMEI,
    start_time,
    end_time=None,
    interpolation=False,
    period=0.1,
):
    """
    Function to get data from IoT dashboard for Local PIM testing
    ```
    import lime_internal_getter as ig
    odf = ig.get_extdata(*("MD0AIOALAA00638", '2024-10-25 14:30', '2024-10-26 02:17'))
    ```
    """
    df = get_data(
        IMEI,
        start_time,
        end_time,
    )
    df["Time diff"] = (
        (pd.to_datetime(df["timeStamp"]).astype("int64") / 10**9).diff().fillna(0)
    )
    df["Cumulative Time"] = df["Time diff"].cumsum().fillna(0)
    col_names = [col for col in df.columns if "Volt" in col and "cell" in col]
    if interpolation:
        time = np.arange(0, df["Cumulative Time"].iloc[-1], period)
        data = [
            pd.Series(time),
            pd.Series(np.interp(time, df["Cumulative Time"], df["batCurrent"])),
        ]
        for col in col_names:
            if not (df[col].eq(0).all()):
                data.append(
                    pd.Series(np.interp(time, df["Cumulative Time"], df[col] / 1000))
                )
    else:
        data = [df["Cumulative Time"], df["batCurrent"]]
        for col in col_names:
            if not (df[col].eq(0).all()):
                data.append(df[col] / 1000)

    return pd.DataFrame(data).T


def get_fwdata(fWVersion="8183D", battery_prefix="MH"):
    """
    Use for getting list of packs with different firmware versions
    import lime_internal_getter as ig
    eg: df= ig.get_fwdata(fWVersion="8183D",battery_prefix="MH")
    """
    headers = headers = {
        "Authorization": f"{auth_key}",
        "Content-Type": "application/json",
    }
    payload = json.dumps(
        {
            "batteryPrefix": battery_prefix,
            "filterParams": [
                {
                    "parameterName": "fwVersion",
                    "downloadType": "Exact",
                    "value": fWVersion,
                }
            ],
            "selectedParams": [
                "socPercent",
                "maxCellSocE",
                "cellSoc3E",
                "minCellSocE",
                "maxCellSohE",
                "minCellSohE",
                "trueSoc",
                "cellSoh3E",
                "socPercentE",
            ],
            "type": "Download",
        }
    )
    response = requests.post(
        "https://api-stage.lime.ai/lime/liveparams/filterDispatchCheckDownloadData",
        headers=headers,
        data=payload,
    )
    return pd.DataFrame(json.loads(response._content)["result"])


def authenticate():
    # Login parameters
    login_params2 = {
        "api": "SYNO.API.Auth",
        "version": "3",
        "method": "login",
        "account": "harisankar",
        "passwd": "8P1GTw#f",
        "session": "FileStation",
        "format": "sid",
    }
    response = requests.get(login_api, params=login_params2)
    data = response.json()
    if data.get("success", False):
        # print("Login successful.")
        return data["data"]["sid"]
    else:
        print("Login failed:", data)
        return None


file_api = "http://10.10.4.40:5000/webapi/entry.cgi"


# Function to read a file via API and process it in memory
def read_file_in_memory(sid, imei, year, month, day, file_date):
    file_path = f"/lime-datalake-prod/lime_bap_parquet/{imei}/{year}/{month}/{day}/{file_date}.parquet"
    file_params = {
        "api": "SYNO.FileStation.Download",
        "version": "2",
        "method": "download",
        "path": file_path,
        "_sid": sid,
    }

    response = requests.get(file_api, params=file_params, stream=True)
    if response.ok:
        try:
            # Read the file into a DataFrame
            file_stream = io.BytesIO(response.content)
            df = pd.read_parquet(file_stream)
            if not isinstance(df, pd.DataFrame):
                raise ValueError("df returned: ", df)
            return df
        except Exception as e:
            raise ValueError(e)
    else:
        raise ValueError(
            "Response: ", response.ok
        )  # Skip printing errors for missing files


def adjust_end_date(end_date: str) -> str:
    """
    Adjusts the end_date if it is greater than or equal to today's date.
    Sets it to yesterday's date in '%d-%m-%Y' format.

    Parameters:
    end_date (str): The end date in '%d-%m-%Y' format.

    Returns:
    str: The adjusted end date in '%d-%m-%Y' format.
    """
    # Parse the input end_date
    end_date_parsed = pd.to_datetime(end_date, format="%Y-%m-%d")
    today = pd.Timestamp.today().normalize()

    # Check if the end_date is >= today
    if end_date_parsed >= today:
        # Set to yesterday
        end_date_parsed = today - pd.Timedelta(days=1)

    # Return the adjusted date in '%d-%m-%Y' format
    return end_date_parsed.strftime("%Y-%m-%d")


def filter_datas(df, start_time, end_time):
    """
    This filters data from start time to end time
    ## Example Usage ##
    df= filter_data(df, '2024-10-25 14:30', '2024-10-26 02:17')

    """
    timestamp = pd.to_datetime(df["date"] + " " + df["time"])
    df = (df[(timestamp >= start_time) & (timestamp <= end_time)]).reset_index(
        drop=True
    )


def get_datas(imei, start_time, end_time=None, filter_data=False, skip=False, nas=True):
    """
    Use for getting battery data from NAS storage eg:
    df = get_data("MD0AIOALAA00638", '2024-10-25 14:30', '2024-10-26 02:17', filter_data=True)
    """
    try:
        int(imei)
        serial_no = False
    except ValueError:
        serial_no = True
    if not end_time:
        end_time = start_time
    if not nas:
        return get_extdata(imei, start_time, end_time, filter_data=filter_data)
    start_date = start_time.split(" ")[0]
    end_date = end_time.split(" ")[0]

    if serial_no:
        try:
            imei = get_imei(imei)
        except Exception:
            raise ValueError("Can't get IMEI")

    end_date = adjust_end_date(end_date)

    sid = authenticate()
    if not sid:
        print("Exiting due to authentication failure.")
        raise ValueError("Authentication Failure")

    # Get dates and ensure they are sorted
    dates = pd.to_datetime(get_dates(start_date, end_date))
    dates = dates.sort_values()

    # Precompute date components to minimize repeated operations
    date_components = [
        (
            d.strftime("%Y"),
            d.strftime("%m").lstrip("0"),
            d.strftime("%d").lstrip("0"),
            d.strftime("%Y%m%d"),
        )
        for d in dates
    ]

    def read_file_wrapper(args):
        """Wrapper function for reading files in parallel."""
        year, month_lstrip, day_lstrip, file_date = args
        try:
            return (
                file_date,
                read_file_in_memory(
                    sid, imei, year, month_lstrip, day_lstrip, file_date
                ),
            )
        except Exception:
            if skip:
                return None
            else:
                raise ValueError(f"Issue getting data for the date {file_date}")

    # Use ThreadPoolExecutor for parallel file reading
    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = [
            executor.submit(read_file_wrapper, components)
            for components in date_components
        ]
        results = []
        for future in concurrent.futures.as_completed(futures):
            try:
                if future.result():
                    results.append(future.result())
            except Exception as e:
                raise e

    # Sort results by file_date to maintain order
    results.sort(key=lambda x: x[0])

    # Extract data frames in the correct order
    data_frames = [result[1] for result in results]

    # Concatenate all data frames in order
    df = pd.concat(data_frames, ignore_index=True)

    if filter_data:
        filter_datas(df, start_time, end_time)

    return df.reset_index(drop=True)


def get_data(imei, start_time, end_time=None, skip=False):
    """
    Use for getting battery data from NAS storage or API
    eg: df= get_data("MD0AIOALAA00638", '2024-10-25 14:30', '2024-10-26 02:17')
    Retrieve device data from NAS storage or via an API as a fallback.

    Args:
        imei (str): The device IMEI to fetch data for.
        start_time (str): The start time in "YYYY-MM-DD HH:mm" format.
        end_time (str, optional): The end time in "YYYY-MM-DD HH:mm" format.
        skip (bool, optional): If True, skip errors. Defaults to False.

    Returns:
        pandas.DataFrame: A DataFrame containing the retrieved data.

    Raises:
        ValueError: If data cannot be retrieved from both NAS and API.
    """
    if not end_time:
        end_time = start_time
    if len(start_time.split(" ")) > 1 or len(end_time.split(" ")) > 1:
        filter_data = True
    else:
        filter_data = False
    try:
        df = get_datas(
            imei,
            start_time,
            end_time=end_time,
            filter_data=filter_data,
            skip=skip,
        )
        df["timeStamp"] = df["date"] + " " + df["time"]
        return df
    except Exception as e:
        try:
            return get_datas(
                imei,
                start_time,
                end_time,
                nas=False,
                filter_data=filter_data,
                skip=skip,
            )
        except Exception as e1:
            raise ValueError("Can't get data from NAS and API: ", e, "\n", e1)


def pim_make(directory_path, model=4, filename="_iot_data.csv"):
    """
    Use this for running the pim after setting configuration inside C code.
    Make sure for temporary testing edit the line no 8 in main.c as
      char *extend = "_iot_data.csv";
      Set appropriate model type in model.h and run pim_make(directory_path,model=4,filename="_iot_data.csv")
    """
    content = f"""
# File name should start with "drive_cycle + FILE_EXT" e.g. drive_cyle_iot_data.csv

FILE_EXT  = {filename}
MODEL_TYPE= {model}

# MODEL_TYPE 1: NMC, 
# MODEL_TYPE 2: LFP, 
# MODEL_TYPE 3: LFP EKF, 
# MODEL_TYPE 4: NMC EKF, 
# MODEL_TYPE 5: NMC VRM, 
# MODEL_TYPE 6: LFP GF, 
# MODEL_TYPE 7: LMFP SINC 
# MODEL_TYPE 8: Gotion LFP
    """
    original_directory = os.getcwd()
    with open(directory_path + "/config.cfg", "w") as f:
        f.write(content)
    try:
        # Change to directory and run make command
        os.chdir(directory_path)
        subprocess.run(["make"], check=True, text=True, capture_output=True)
    except subprocess.CalledProcessError as e:
        print(f"Error running make: {e.stderr}")
    finally:
        # Change back to the original directory
        os.chdir(original_directory)


def soh_calculator(df, capacity, x3, y3):
    time_diff = df["Cumulative Time"].diff().fillna(0)
    capacity = capacity  # Ah
    soh = 1.0
    timer = 0
    cumu_ah = 0
    prev_soc = np.interp(df["Voltage(V)"].iloc[0], y3, x3)
    soc1 = prev_soc
    for o in range(1, len(df)):
        if df["Step Type"].iloc[o] == "Rest":
            timer += time_diff.iloc[o]
        else:
            timer = 0
        soc_diff = (df["Current(A)"].iloc[o] * time_diff.iloc[o] / 3600) / (
            capacity * soh
        )
        soc1 += soc_diff
        cumu_ah += df["Current(A)"].iloc[o] * time_diff.iloc[o] / 3600
        if timer > 900 and abs(cumu_ah / capacity) > 0.5:
            soc2 = np.interp(df["Voltage(V)"].iloc[o], y3, x3)
            soh = abs(1 / ((((soc1 - soc2) * capacity) / (cumu_ah)) - (1 / soh)))
            break
    return soh


def chromasoc_calculator(df, capacity, x3, y3, soh, initial_soc=None):
    if initial_soc:
        chromasoc = [initial_soc]
    else:
        chromasoc = [np.interp(df["Voltage(V)"].iloc[0], y3, x3)]
    time_diff = df["Time diff"]
    timer = 0
    for o in range(1, len(df)):
        if df["Step Type"].iloc[o] == "Rest":
            timer += time_diff.iloc[o]
        else:
            timer = 0
        soc_diff = (df["Current(A)"].iloc[o] * time_diff.iloc[o] / 3600) / (
            capacity * soh
        )
        if timer > 900:
            chromasoc.append(np.interp(df["Voltage(V)"].iloc[o], y3, x3))
            timer = 0
        else:
            chromasoc.append(chromasoc[-1] + soc_diff)
    return pd.Series(chromasoc) * 100


def time_convert(x):
    """
    Converts a time string in the format "HH:MM:SS" to seconds

    Parameters
    ----------
    x : str
        The time string to convert

    Returns
    -------
    int
        The time in seconds
    """
    x1 = list(map(int, x.split(":")))
    return x1[0] * 3600 + x1[1] * 60 + x1[2]


def derive_etair(data, capacity, x3, y3, plot=False):
    """
    Derives the eta-IR model parameters from a given data set

    Parameters
    ----------
    data : pandas.DataFrame
        The data to derive the model from. Must contain columns 'Total Time', 'Current(A)', 'Step Type', 'Step Index', and 'Voltage(V)'
    capacity : float
        The capacity of the battery in Ah
    x3 : numpy.ndarray
        The SOC array for the OCV curve
    y3 : numpy.ndarray
        The OCV array for the OCV curve
    plot : bool, optional
        If True, plots the optimisation results for each discharge rate. Defaults to False

    Returns
    -------
    tuple
        A tuple containing the derived eta and internal resistance values
    """
    if isinstance(data, str):
        data = pd.read_excel(data, sheet_name="record")
    df = ((data[data["Step Index"] > 1].iloc[50:]).reset_index(drop=True)).copy()
    df["Cumulative Time"] = df["Total Time"].apply(time_convert)
    df["Cumulative Time"] = df["Cumulative Time"].diff().fillna(0).cumsum()
    df["Time diff"] = df["Cumulative Time"].diff().fillna(0)
    capacity = capacity  # Ah
    soh = soh_calculator(df, capacity, x3, y3)
    print("Derived SOH: ", soh)
    df["ChromaSOC"] = chromasoc_calculator(df, capacity, x3, y3, soh)
    current = df["Current(A)"]
    current_df = df[current < 0].copy()
    soc_rates = (current_df["Current(A)"] / 15).round(decimals=2).unique()
    current_df["soc_rates"] = (current_df["Current(A)"] / 15).round(decimals=2)
    x3_dv = np.arange(0, 1.0, 0.01)
    y3_dv = np.interp(x3_dv, x3, y3)
    y3_diff = np.diff(y3_dv, prepend=1.0) / np.diff(x3_dv, prepend=1.0)
    eta_ir = []
    j0 = []
    for s in soc_rates:
        soc_df = current_df[current_df["soc_rates"] == s].copy()
        soc_df["dv"] = np.interp(soc_df["ChromaSOC"] / 100, x3_dv, y3_diff)
        derive_df = soc_df[soc_df["dv"] > 0.5]

        def error(params):
            global soc_noise
            a, b = abs(params)
            r1 = a * 0.001 * ((derive_df["Current(A)"] / 1) / (capacity * soh))
            r2 = ((2 * 8.314462618 * (298.15)) / 96485.33212) * np.arcsinh(
                (derive_df["Current(A)"] / 1) / (2 * (1 / b) * (capacity * soh))
            )
            soc_noise = np.interp(derive_df["Voltage(V)"] - (r1 + r2), y3, x3)
            soc_noise = [o * 100 for o in soc_noise]
            return (
                np.sqrt(
                    mean_squared_error(pd.Series(soc_noise), derive_df["ChromaSOC"])
                )
                * 1000
            )

        a_guess = 10
        b_guess = 0.1
        res = scipy.optimize.minimize(
            error,
            [
                a_guess,
                b_guess,
            ],  # ,bounds=[(1.61880320e+02,2.12495519e+02),(9.42777888e-02,1.04222729e-01)],
            method="Powell",  # optimising for process noises and covariances
        )
        eta_ir.append(res.x[0])
        j0.append(res.x[1])
        if plot:
            print(res.x, end=" ")
            print(error(res.x) / 1000)
            fig = go.Figure()
            fig.add_trace(
                go.Scatter(
                    x=derive_df["Cumulative Time"],
                    y=soc_noise,
                    name="optimised R",
                    line=dict(dash="solid"),
                )
            )
            fig.add_trace(
                go.Scatter(
                    x=derive_df["Cumulative Time"],
                    y=derive_df["ChromaSOC"],
                    name="Chroma",
                    line=dict(dash="dot"),
                    mode="markers",
                )
            )
            fig.update_layout(dict(title=f"optimisation for c_rate {s}"))
            fig.show(renderer="browser")
    if plot:
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=abs(soc_rates), y=eta_ir, mode="markers"))
        fig.update_layout(
            title="Derived eta IR", xaxis_title="C-rates", yaxis_title="Eta-IR"
        )
        fig.show(renderer="browser")
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=abs(soc_rates), y=j0, mode="markers"))
        fig.update_layout(
            title="Derived invJ0", xaxis_title="C-rates", yaxis_title="inv J0"
        )
        fig.show(renderer="browser")
    return (eta_ir, j0)


class PIMProcessor:
    """
    Use this for processing the model and generating tables for lists of batteries.
    This version always processes both SOC and SOH without any step argument.
    """

    def __init__(self, directory_path=None, model=4):
        if directory_path:
            self.directory_path = directory_path
        else:
            self.directory_path = None
        self.model = model
        self.fdf = None
        self.final_table = None
        self.__odf_soc = None
        self.__odf_soh = None
        if self.model in [3, 6, 8, 9]:
            self.__lfp_model = True
        else:
            self.__lfp_model = False

    def fetch_and_process_data(
        self,
        serial_numbers,
        start_date,
        end_date=None,
        soh_value=None,
        interpolation=False,
        error_display=True,
        fw=None,
    ):
        """
        Fetches and processes data for each serial number and appends SOC and SOH
        (and their min/max) to the DataFrame.
        Final data is stored in self.fdf.
        """
        if not end_date:
            end_date = start_date
        k = 0
        for ser in tqdm(serial_numbers, desc="Processing Serial Numbers", leave=False):
            params = (ser, start_date, end_date)
            if self.directory_path:
                pim_df = get_pimdata(
                    params[0], params[1], params[2], interpolation=interpolation
                )
                # 1) Fetch IoT data and save as CSV
                try:
                    pim_df.to_csv(
                        f"{self.directory_path}/input_data/drive_cycle_iot_data.csv",
                        index=False,
                        header=None,
                    )
                except Exception as e:
                    print(f"Error fetching IoT data for {ser}: {e}")
                    continue

                # 2) Run the make command to generate model outputs
                try:
                    pim_make(self.directory_path, model=self.model)
                except subprocess.CalledProcessError as e:
                    print(f"Error running make for {ser}: {e.stderr}")
                    continue

            # 3) Get the combined DataFrame from internal storage
            try:
                df = get_data(params[0], params[1], params[2])
                try:
                    if fw:
                        if any(
                            firmware not in fw for firmware in df.fwVersion.unique()
                        ):
                            raise ValueError("Firmware version not detected")
                except Exception as e:
                    if error_display:
                        print(f"Skipping {ser} due to firmware issue: {e}")
                    continue
            except Exception as e:
                if error_display:
                    print(f"Error fetching extended data for {ser}: {e}")
                continue

            df["Time diff"] = (
                (pd.to_datetime(df["timeStamp"]).astype("int64") / 10**9)
                .diff()
                .fillna(0)
            )
            df["Cumulative Time"] = df["Time diff"].cumsum().fillna(0)
            time_data = pd.to_datetime(df["timeStamp"]).astype("int64") / 10**9
            if self.directory_path:
                # 4) Read final_soc_iot_data.csv
                try:
                    odf_soc = pd.read_csv(
                        f"{self.directory_path}/output_data/final_soc_iot_data.csv",
                        header=None,
                    )
                    self.__odf_soc = odf_soc
                    maxx_soc = odf_soc.iloc[:, 1:].max(axis=1)
                    minn_soc = odf_soc.iloc[:, 1:].min(axis=1)
                    df["PIM_Soc"] = np.interp(
                        df["Cumulative Time"],
                        odf_soc[0].shift().fillna(0),
                        odf_soc[3],
                    )
                    df["PIM_maxSoc"] = np.interp(
                        df["Cumulative Time"],
                        odf_soc[0].shift().fillna(0),
                        maxx_soc,
                    )
                    df["PIM_minSoc"] = np.interp(
                        df["Cumulative Time"],
                        odf_soc[0].shift().fillna(0),
                        minn_soc,
                    )
                except Exception as e:
                    print(f"Error reading final_soc_iot_data.csv for {ser}: {e}")

                # 5) Read final_soh_iot_data.csv
                try:
                    odf_soh = pd.read_csv(
                        f"{self.directory_path}/output_data/final_soh_iot_data.csv",
                        header=None,
                    )
                    self.__odf_soh = odf_soh
                    maxx_soh = odf_soh.iloc[:, 1:].max(axis=1)
                    minn_soh = odf_soh.iloc[:, 1:].min(axis=1)
                    df["PIM_Soh"] = np.interp(
                        df["Cumulative Time"],
                        odf_soh[0].shift().fillna(0),
                        odf_soh[3],
                    )
                    df["PIM_maxSoh"] = np.interp(
                        df["Cumulative Time"],
                        odf_soh[0].shift().fillna(0),
                        maxx_soh,
                    )
                    df["PIM_minSoh"] = np.interp(
                        df["Cumulative Time"],
                        odf_soh[0].shift().fillna(0),
                        minn_soh,
                    )
                except Exception as e:
                    print(f"Error reading final_soh_iot_data.csv for {ser}: {e}")
            # Convert back to real timestamps and keep track of serial number
            df["Cumulative Time"] = pd.to_datetime(time_data, unit="s")
            df["Serial_no"] = ser
            # Python Model runner
            if self.model == 9:
                kf = KalmanFilter(self.model)
                kf.process_filter(
                    pim_df[1:].reset_index(drop=True), soh_value=soh_value
                )
                df["Python Soc"] = kf.soc[2] * 100
                kfd = pd.DataFrame(kf.soc).T
                df["Python Max Soc"] = (kfd.max(axis=1)) * 100
                df["Python Min Soc"] = (kfd.min(axis=1)) * 100
                df["Python Soh"] = kf.soh[2] * 100
                kfh = pd.DataFrame(kf.soh).T
                df["Python Max Soh"] = (kfh.max(axis=1)) * 100
                df["Python Min Soh"] = (kfh.min(axis=1)) * 100
                df["Used Measurements Soc"] = kf.used_measurements[2] * 100
                df["Measurements Soc"] = kf.measurements[2] * 100

            if k > 0:
                self.fdf = pd.concat([self.fdf, df.copy()])
            else:
                self.fdf = df.copy()
                k += 1

    def generate_final_table(self, save_csv=False):
        """
        Generates a final table containing both SOC and SOH details from the data.
        Final table is stored in self.final_table.
        """
        if self.fdf is None:
            print("No data available. Please run fetch_and_process_data first.")
            return

        self.final_table = None
        k = 0
        for ser in tqdm(
            self.fdf["Serial_no"].unique(),
            desc="Generating Final Table",
            leave=False,
        ):
            table = self.fdf[self.fdf["Serial_no"] == ser][:-2:-1].copy()
            if self.directory_path:
                # Scale PIM columns by 100
                for col in [col for col in table.columns if "PIM_" in col]:
                    table[col] *= 100

            # Keep relevant SOC and SOH columns
            soc_cols = [
                col
                for col in table.columns
                if "Soc" in col
                and "Raw" not in col
                and "reserve" not in col
                and "Warn" not in col
            ]
            soh_cols = [
                col
                for col in table.columns
                if "Soh" in col
                and "Raw" not in col
                and "reserve" not in col
                and "Warn" not in col
            ]
            needed_cols = soc_cols + soh_cols

            if k == 0:
                self.final_table = table[needed_cols]
                k += 1
            else:
                self.final_table = pd.concat([self.final_table, table[needed_cols]])

        if save_csv and self.final_table is not None:
            self.final_table.to_csv(
                f"{self.directory_path}/final_table.csv", index=False
            )

    def plot(self, renderer="browser", plot_all=False):
        """
        Plots two separate windows for SOC and SOH for each serial number
        in the data by default.
        Parameters:
        - renderer: Use None to display the plot in the console.
        - plot_all: If True, plots all available columns for SOC and SOH.
        """
        if self.fdf is None:
            print("No data to plot. Please run fetch_and_process_data first.")
            return

        # For each serial number, create two plots: SOC and SOH
        for ser in self.fdf["Serial_no"].unique():
            table = self.fdf[self.fdf["Serial_no"] == ser].copy()

            # Scale PIM columns by 100
            for col in [col for col in table.columns if "PIM_" in col]:
                table[col] *= 100

            # 1) SOC Plot
            soc_cols = [
                col
                for col in table.columns
                if "Soc" in col
                and "Raw" not in col
                and "reserve" not in col
                and "Warn" not in col
            ]
            fig_soc = go.Figure(layout=dict(title=f"SOC Comparison for {ser}"))

            if plot_all:
                for tcol in soc_cols:
                    fig_soc.add_trace(
                        go.Scatter(
                            x=table["Cumulative Time"],
                            y=table[tcol],
                            name=tcol,
                        )
                    )
            else:
                if self.directory_path:
                    # Common columns to plot (if present)
                    default_soc_cols = ["trueSoc", "cellSoc3E", "PIM_Soc"]
                else:
                    default_soc_cols = ["trueSoc", "cellSoc3E"]
                for tcol in default_soc_cols:
                    if tcol in table.columns:
                        fig_soc.add_trace(
                            go.Scatter(
                                x=table["Cumulative Time"],
                                y=table[tcol],
                                name=tcol,
                            )
                        )
            fig_soc.show(renderer=renderer)

            # 2) SOH Plot
            soh_cols = [
                col
                for col in table.columns
                if "Soh" in col
                and "Raw" not in col
                and "reserve" not in col
                and "Warn" not in col
            ]
            fig_soh = go.Figure(layout=dict(title=f"SOH Comparison for {ser}"))

            if plot_all:
                for tcol in soh_cols:
                    fig_soh.add_trace(
                        go.Scatter(
                            x=table["Cumulative Time"],
                            y=table[tcol],
                            name=tcol,
                        )
                    )
            else:
                # Common columns to plot (if present)
                if self.directory_path:
                    default_soh_cols = ["sohPercent", "cellSoh3E", "PIM_Soh"]
                else:
                    default_soh_cols = ["sohPercent", "cellSoh3E"]
                for tcol in default_soh_cols:
                    if tcol in table.columns:
                        fig_soh.add_trace(
                            go.Scatter(
                                x=table["Cumulative Time"],
                                y=table[tcol],
                                name=tcol,
                            )
                        )
            fig_soh.show(renderer=renderer)

    def __correction_conditions(self, soc1, soc2, soc3=None):
        if self.__lfp_model:
            if soc3:
                if (soc1 < 20) and (soc2 < 20) and (soc3 < 20):
                    return True
            else:
                if (soc1 < 20) and (soc2 < 20):
                    return True
            return False
        else:
            return True

    def __process_row_for_errors(self, i, devi=300):
        global df1, odf
        serial_num = i
        soc_cat = np.arange(0, 100, 5)
        if self.directory_path:
            if self.__odf_soc is not None:
                odf = self.__odf_soc
            else:
                raise RuntimeError("Please run fetch_and_process data first")
        try:
            # Fetch data for max_error calculation
            df1 = self.fdf[self.fdf["Serial_no"] == serial_num].copy()
            if df1.shape[0] == 0:
                return serial_num, None, None
        except KeyError:
            print(f"Skipping for {serial_num} for error calculation due data error")
            return serial_num, None, None
        df1["Time diff"] = (
            (pd.to_datetime(df1["timeStamp"]).astype("int64") / 10**9).diff().fillna(0)
        )
        time_data = pd.to_datetime(df1["timeStamp"]).astype("int64") / 10**9
        if self.directory_path:
            time_datao = odf[0].diff()
            time_datao.iloc[0] = time_data.iloc[0]
            df1["PIM_SOC"] = np.interp(
                df1["Time diff"].cumsum().fillna(0), odf[0], odf[3]
            )
            time_datao = time_datao.cumsum().fillna(
                time_datao.iloc[0]
            )  # Convert output time to the format of input timestamp
            time_datao = pd.to_datetime(time_datao, unit="s")
        time_data = pd.to_datetime(time_data, unit="s")
        timer = 0
        error1 = []
        time1 = []
        error2 = []
        soc_range = []
        dev = []
        for o in range(df1.shape[0] - devi):
            if abs(df1["batCurrent"].iloc[o]) < 0.5:
                timer += df1["Time diff"].iloc[o]
            else:
                timer = 0
            if timer > 900 and (
                ((df1["trueSoc"].iloc[o] - df1["trueSoc"].iloc[o + 10]) > 0.05)
                and (abs(df1["batCurrent"].iloc[o + 10]) < 0.5)
            ):
                if self.directory_path:
                    if self.__correction_conditions(
                        df1["trueSoc"].iloc[o + 10],
                        df1["cellSoc3E"].iloc[o + 10],
                        df1["PIM_SOC"].iloc[o + 10],
                    ) and (
                        df1["trueSoc"].iloc[o + 10] != 0
                        and df1["cellSoc3E"].iloc[o + 10] != 0
                        and df1["trueSoc"].iloc[o] != 0
                        and df1["PIM_SOC"].iloc[o] != 0
                        and df1["PIM_SOC"].iloc[o + 10] != 0
                    ):
                        for soc_c in soc_cat:
                            if df1["trueSoc"].iloc[o + 10] <= soc_c:
                                soc_range.append(int(soc_c) + 10)
                                error1.append(
                                    df1["trueSoc"].iloc[o + 10]
                                    - df1["cellSoc3E"].iloc[o]
                                )
                                error2.append(
                                    df1["trueSoc"].iloc[o + 10]
                                    - (df1["PIM_SOC"] * 100).iloc[o]
                                )
                                time1.append(df1["timeStamp"].iloc[o])
                                time_dev = (
                                    pd.to_datetime(
                                        df1["timeStamp"].iloc[o + devi]
                                    ).value
                                    / (10**9)
                                ) - (pd.to_datetime(time1[-1]).value / (10**9))
                                dev.append(
                                    (
                                        df1["trueSoc"].iloc[o + devi]
                                        - df1["cellSoc3E"].iloc[o + devi]
                                    )
                                    / time_dev
                                )
                                break

                else:
                    if self.__correction_conditions(
                        df1["trueSoc"].iloc[o + 10],
                        df1["cellSoc3E"].iloc[o + 10],
                    ) and (
                        df1["trueSoc"].iloc[o + 10] != 0
                        and df1["cellSoc3E"].iloc[o + 10] != 0
                        and df1["trueSoc"].iloc[o] != 0
                    ):
                        for soc_c in soc_cat:
                            if df1["trueSoc"].iloc[o + 10] <= soc_c:
                                soc_range.append(int(soc_c))
                                error1.append(
                                    df1["trueSoc"].iloc[o + 10]
                                    - df1["cellSoc3E"].iloc[o]
                                )
                                time1.append(df1["timeStamp"].iloc[o])
                                time_dev = (
                                    pd.to_datetime(
                                        df1["timeStamp"].iloc[o + devi]
                                    ).value
                                    / (10**9)
                                ) - (pd.to_datetime(time1[-1]).value / (10**9))
                                dev.append(
                                    (
                                        df1["trueSoc"].iloc[o + devi]
                                        - df1["cellSoc3E"].iloc[o + devi]
                                    )
                                    / time_dev
                                )
                                break

        return serial_num, error1, error2, time1, soc_range, dev

    def correction_monitor(self, renderer="browser", plot=True, deviation=200):
        """
        Use for plotting
        """
        self.bms_error = []
        self.pim_error = []
        self.correction_time = []
        self.serial_num = []
        self.soc_range = []
        self.deviation = []
        for ser in tqdm(
            self.fdf["Serial_no"].unique(),
            desc="Calculating corrections for serial numbers",
            leave=False,
        ):
            result = self.__process_row_for_errors(ser, devi=deviation)
            result = list(result)
            if result[1]:
                self.bms_error.append(abs(pd.Series(result[1])))
                self.pim_error.append(abs(pd.Series(result[2])))
                self.correction_time.append(pd.Series(result[3]))
                self.serial_num.append(ser)
                self.soc_range.append(pd.Series(result[4]))
                self.deviation.append(pd.Series(result[5]))
            if plot and result[1]:
                fig = go.Figure(
                    layout=dict(
                        title="Error comparison during SOC correction in bms and PIM for "
                        + str(ser)
                    )
                )
                fig.add_trace(
                    go.Scatter(
                        x=pd.to_datetime(result[3]),
                        y=abs(pd.Series(result[1])),
                        name="BMS SOH :" + str(self.fdf["cellSoh3E"].iloc[-1]),
                    )
                )
                if self.directory_path:
                    fig.add_trace(
                        go.Scatter(
                            x=pd.to_datetime(result[3]),
                            y=abs(pd.Series(result[2])),
                            name="PIM SOH :"
                            + str(
                                np.round(
                                    self.fdf["PIM_Soh"].iloc[-1] * 100,
                                    decimals=0,
                                )
                            ),
                        )
                    )
                if renderer:
                    fig.show(renderer="browser")
                else:
                    fig.show()


class Reports:
    """
    A class to generate reports.
    Methods
    -------
    df_to_docx(data, heading="Table data", save=False, file_name="Table data.docx")
        Converts a DataFrame to a Word document with a table.
    correction_script(fWlist, model, start_date, end_date=None, batch_size=30, prefix="", script="soc_range")

    """

    def __init__(self):
        self.document = Document()

    def df_to_docx(
        self,
        data,
        heading="Table data",
        save=False,
        file_name="Table data.docx",
    ):
        doc = self.document
        df = data
        doc.add_heading(heading, level=1)
        # Add a table to the document
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"

        # Add header row
        header_cells = table.rows[0].cells
        for idx, col_name in enumerate(df.columns):
            header_cells[idx].text = col_name

        # Add data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                row_cells[idx].text = str(np.round(value, decimals=4))
        if save:
            doc.save(file_name)

    def correction_script(
        self,
        fWlist,
        model,
        start_date,
        end_date=None,
        batch_size=30,
        prefix="",
        script="all",
    ):
        """
        def correction_script(fWlist, model, start_date, end_date=None, batch_size=30, prefix="", script='soc_range'):

            Processes firmware data and generates correction reports.
            Parameters:
            fWlist (list): List of firmware versions to process.
            model (object): Model object used for processing.
            start_date (str): Start date for data processing.
            end_date (str, optional): End date for data processing. Defaults to None.
            batch_size (int, optional): Number of records to process in each batch. Defaults to 30.
            prefix (str, optional): Prefix for battery data. Defaults to "".
            script (str, optional): Type of script to run. Options are 'soc_range', 'max_error', or 'all'. Defaults to 'soc_range'.
            Returns:
            None
            This function processes the firmware data for each firmware version in fWlist. It fetches and processes data in batches,
            and generates correction reports based on the specified script type. The reports are saved as CSV files.
        """

        for fW in fWlist:
            kdf = get_fwdata(fWVersion=fW, battery_prefix=prefix)
            processor = PIMProcessor(model=model)
            batch_size = batch_size
            filepath = f"{fW}" + "soc_range" + ".csv"
            filepath2 = f"{fW}" + ".csv"

            sc = []
            be = []
            ct = []
            ser = []
            dev = []

            for batch in tqdm(
                range(0, kdf.shape[0], batch_size),
                desc=f"Processing batch for {fW}",
            ):
                processor.fetch_and_process_data(
                    kdf["imei"][batch : batch + batch_size],
                    start_date=start_date,
                    end_date=end_date,
                    error_display=False,
                )
                processor.correction_monitor(plot=False)
                if (script == "soc_range") or (script == "all"):
                    for i in range(len(processor.soc_range)):
                        sc.extend(processor.soc_range[i])
                        be.extend(processor.bms_error[i])
                        ct.extend(processor.correction_time[i])
                        dev.extend(processor.deviation[i])
                        ser.extend(
                            [
                                processor.serial_num[i]
                                for o in range(len(processor.soc_range[i]))
                            ]
                        )
                    df = pd.DataFrame(
                        {
                            "imei": ser,
                            "bms_error": be,
                            "correction_time": ct,
                            "soc_range": sc,
                            "deviation": dev,
                        }
                    )
                    if not os.path.exists(filepath):
                        df.to_csv(filepath, index=False)
                    else:
                        df.to_csv(filepath, mode="a", header=False, index=False)
                elif (script == "all") or (script == "max_error"):
                    max_error = pd.Series([i.max() for i in processor.bms_error])
                    min_error = pd.Series([i.min() for i in processor.bms_error])
                    average_error = pd.Series([i.mean() for i in processor.bms_error])
                    imei = pd.Series(processor.serial_num)
                    df = pd.DataFrame(
                        {
                            "imei": imei,
                            "max_error": max_error,
                            "min_error": min_error,
                            "average_error": average_error,
                        }
                    )
                    if not os.path.exists(filepath2):
                        df.to_csv(filepath2, index=False)
                    else:
                        df.to_csv(filepath2, mode="a", header=False, index=False)
