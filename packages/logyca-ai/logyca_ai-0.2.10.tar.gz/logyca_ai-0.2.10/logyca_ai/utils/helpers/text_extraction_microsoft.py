from docx import Document # python-docx
from enum import StrEnum
from io import BytesIO
from logyca_ai.utils.constants.ocr import OCREngine, OCREngineSettings
from logyca_ai.utils.helpers.garbage_collector_helper import garbage_collector_at_the_end
from logyca_ai.utils.schemes.output.conversations import ImageBase64
from openpyxl import load_workbook
from PIL import Image  # Pillow
import base64
import json
import os
import pandas as pd
import pytesseract

class ExcelSheetStates(StrEnum):
    HIDDEN = "hidden"
    VERY_HIDDEN = "veryHidden"
    VISIBLE = "visible"

@garbage_collector_at_the_end
def extract_text_from_docx_file(filename_full_path:str|BytesIO,advanced_image_recognition:bool=False,ocr_engine_path:str=None,output_temp_dir:str=None):
    """
    Extracts text from a DOCX file.

    :param filename_full_path: Full path to the DOCX file from which to extract text or file loaded in BytesIO RAM memory.
    :type filename_full_path: str|BytesIO
    :param advanced_image_recognition: Indicates whether to perform text recognition on images within the DOCX.
                               If True, OCR techniques will be used to extract text from images.
    :type advanced_image_recognition: bool
    :param ocr_engine_path: Path to the OCR executable. If provided, this path will be used instead of the default.
    :type ocr_engine_path: str, optional
    :param output_temp_dir: Temporary directory for storing output files.
                            If not provided, a default tmp temporary directory in the application root folder will be used.
    :type output_temp_dir: str, optional

    :return: Extracted text from the DOCX file.
    :rtype: str

    :raises FileNotFoundError: If the specified DOCX file is not found.
    :raises ValueError: If the OCR path is invalid.
    """

    if advanced_image_recognition:
        if ocr_engine_path is None:
            pytesseract.pytesseract.tesseract_cmd=OCREngine.get_binary_path()
        else:
            pytesseract.pytesseract.tesseract_cmd=ocr_engine_path

    doc = Document(filename_full_path)

    text = ""
    if output_temp_dir is None:
        output_temp_dir=os.path.abspath(os.path.join(os.path.dirname(__file__),OCREngineSettings.TMP_DIR))
    if not os.path.exists(output_temp_dir):
        os.makedirs(output_temp_dir)

    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    if advanced_image_recognition:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob

                    image = Image.open(BytesIO(image_bytes))
                    ocr_text = pytesseract.image_to_string(image)
                    text += ocr_text
                except:
                    """If the image format is not supported, the image will be skipped."""
                    pass

    return text

@garbage_collector_at_the_end
def extract_images_from_docx_file(filename_full_path:str|BytesIO)->list:
    """
    Extracts text from a DOCX file.

    :param filename_full_path: Full path to the DOCX file from which to extract text or file loaded in BytesIO RAM memory.
    :type filename_full_path: str|BytesIO

    :return: Extracted text from the DOCX file.
    :rtype: list

    :raises FileNotFoundError: If the specified DOCX file is not found.
    :raises ValueError: If the OCR path is invalid.
    """

    doc = Document(filename_full_path)
    images = []

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_part = rel.target_part
                image_bytes = image_part.blob

                img = Image.open(BytesIO(image_bytes))
                image_format = img.format if img.format else "PNG"
                image_format = str(image_format).lower()
                buffered = BytesIO()
                img.save(buffered, format=image_format)
                image_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
                images.append(ImageBase64(
                    image_base64=image_base64,
                    image_format=image_format
                    ).to_dict()
                )
                del buffered
            except:
                """If the image format is not supported, the image will be skipped."""
                pass

    return images

@garbage_collector_at_the_end
def extract_text_from_excel_file(filename_full_path: str|BytesIO, advanced_image_recognition: bool = False, ocr_engine_path: str = None, output_temp_dir: str = None,format_output:str="json",only_sheet_visibility_enabled:bool=True):
    """
    Extracts text from an Excel file including all sheets and any embedded images.

    :param filename_full_path: Full path to the Excel file from which to extract text or file loaded in BytesIO RAM memory.
    :param advanced_image_recognition: Indicates whether to perform OCR on images within the Excel file.
    :param ocr_engine_path: Path to the OCR executable.
    :param output_temp_dir: Temporary directory for storing output files.
    :param format_output: Data array format for output, can be csv, json, list. The default is json.
    :param only_sheet_visibility_enabled: When looping through excel sheets, filter those with visible status if true, if false, do not loop through sheets with another status.
    :return: A text in JSON format representing the structure of a Microsoft Excel spreadsheet and has the following format.
                Sheet_name: It is the name of the primary key that represents an object, within this object, there are several subkeys.
                Sheet_name: It is the name of the primary key.
                content: Contains a list of lists (a two-dimensional array) that represents the content of the spreadsheet cells.
                image_text: Contains the text extracted from images within the Microsoft Excel document.
                Example of the content in the list of formats that I will send you.
                {
                    "sheet_name": {
                        "content": [["", "", "", ""],["", "", "", ""]],
                        "image_text": ""
                        },
                }
    :rtype: dict
    """
    # Set OCR engine path
    if advanced_image_recognition:
        if ocr_engine_path is None:
            pytesseract.pytesseract.tesseract_cmd = OCREngine.get_binary_path()
        else:
            pytesseract.pytesseract.tesseract_cmd = ocr_engine_path

    if output_temp_dir is None:
        output_temp_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), OCREngineSettings.TMP_DIR))
    if not os.path.exists(output_temp_dir):
        os.makedirs(output_temp_dir)

    result = {}
    sheets_description = {}

    ################################################
    # OCR - Extract data from images

    workbook = load_workbook(filename_full_path, data_only=True)
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        if only_sheet_visibility_enabled is False:
            make_extraction = True
        elif sheet.sheet_state == ExcelSheetStates.VISIBLE:
            make_extraction = True
        else:
            make_extraction = False

        if make_extraction is True:
            text = ""
            if advanced_image_recognition:
                for image in sheet._images:
                    try:
                        img = Image.open(BytesIO(image._data()))
                        ocr_text = pytesseract.image_to_string(img)
                        text += ocr_text
                        text += "\n"
                    except:
                        """If the image format is not supported, the image will be skipped."""
                        pass                    
            result[sheet_name] = {
                "image_text": text
            }
            sheets_description[sheet_name] = {
                "visibility": sheet.sheet_state
            }
            del sheet
    del workbook

    ################################################
    # Read text data
    excel_data = pd.ExcelFile(filename_full_path)
    # print(f"memory get_memory_consumed_by_current_process:{get_memory_consumed_by_current_process()}")
    sheet_names = excel_data.sheet_names
    for sheet in sheet_names:

        if sheets_description.get(sheet,False):
            make_extraction = True
        else:
            make_extraction = False
        if make_extraction is True:
            df = pd.read_excel(filename_full_path, sheet_name=sheet, header=None)
            df = df.fillna("")
            df = df.dropna(how='all')

            if format_output.lower() == 'list':
                list_of_lists = df.values.tolist()
            elif format_output.lower() == 'csv':
                list_of_lists = df.to_csv(index=False, sep=";", lineterminator='\n')
            elif format_output.lower() == 'json':
                list_of_lists = df.to_json(orient='records')
            else:
                raise ValueError(f"Unsupported format format_output: {format_output}")

            # print(list_of_lists)
            # print(df)
            result[sheet]["content"]=list_of_lists
            del df
            del list_of_lists

    excel_data.close()
    del excel_data

    ################################################
    # Finish

    json_result = json.dumps(result,default=str)
    # print(json_result)
    return json_result

@garbage_collector_at_the_end
def extract_images_excel_file(filename_full_path: str|BytesIO,only_sheet_visibility_enabled:bool=True)->list:
    """
    Extracts text from an Excel file including all sheets and any embedded images.

    :param filename_full_path: Full path to the Excel file from which to extract text or file loaded in BytesIO RAM memory.
    :rtype: list
    :param only_sheet_visibility_enabled: When looping through excel sheets, filter those with visible status if true, if false, do not loop through sheets with another status.
    :rtype: bool
    """

    workbook = load_workbook(filename_full_path, data_only=True)
    images = []
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        
        if only_sheet_visibility_enabled is False:
            make_extraction = True
        elif sheet.sheet_state == ExcelSheetStates.VISIBLE:
            make_extraction = True
        else:
            make_extraction = False
        
        if make_extraction is True:
            for image in sheet._images:
                try:
                    img = Image.open(BytesIO(image._data()))
                    image_format = img.format if img.format else "PNG"
                    image_format = str(image_format).lower()
                    buffered = BytesIO()
                    img.save(buffered, format=image_format)
                    image_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
                    images.append(ImageBase64(
                        image_base64=image_base64,
                        image_format=image_format,
                        sheet_name=sheet_name
                        ).to_dict()
                    )
                    del buffered
                except:
                    """If the image format is not supported, the image will be skipped."""
                    pass
        del sheet
    del workbook

    return images
