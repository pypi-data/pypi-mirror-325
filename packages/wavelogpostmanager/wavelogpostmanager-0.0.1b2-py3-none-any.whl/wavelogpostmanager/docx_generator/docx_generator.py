#!/usr/bin/env python3.10
# -*- coding: utf-8 -*-
# author： NearlyHeadlessJack
# email: wang@rjack.cn
# datetime： 2025/1/28 17:47
# ide： PyCharm
# file: docx_generator.py
"""
input: the dict list of qso [{zip_code, address, name, callsign, sign-off code, phone number},{},{}]

create one docx of all envelopes
"""
import shutil
from docx.shared import Inches
from docx import Document
import os
import glob
import datetime
import qrcode
from docxcompose.composer import Composer
from wavelogpostmanager.constants.languages import Language as L
from wavelogpostmanager.config import ConfigContext


class DocxGenerator:
    # default generator folder
    path = r"./docx"
    # hidden folder of docx
    mid_folder = r".docx"
    # hidden folder of img
    img_folder = r".img"
    # qr generate config
    qr_context = {
        "version": 1,
        "error_correction": qrcode.constants.ERROR_CORRECT_H,
        "box_size": 10,
        "border": 1,
    }

    @staticmethod
    def _init():
        if not os.path.exists(DocxGenerator.path):
            os.mkdir(DocxGenerator.path)
        if not os.path.exists(DocxGenerator.path + "/" + DocxGenerator.mid_folder):
            os.mkdir(DocxGenerator.path + "/" + DocxGenerator.mid_folder)
        if not os.path.exists(DocxGenerator.path + "/" + DocxGenerator.img_folder):
            os.mkdir(DocxGenerator.path + "/" + DocxGenerator.img_folder)

    @staticmethod
    def _check_template() -> int:
        template_path = ConfigContext.config["docx_generator"]["template_path"]
        if not os.path.exists(template_path):
            print(f"-{L.get('template_not_found')}")
            print(f"-{L.get('template_not_found_hint')}{DocxGenerator.path}")
            return 1
        return 0

    @staticmethod
    def _generator_all_envelopes(qso_list: list) -> int:
        DocxGenerator._init()
        if DocxGenerator._check_template():
            return 1

        for idx, qso in enumerate(qso_list):
            template_path = ConfigContext.config["docx_generator"]["template_path"]
            doc = Document(template_path)
            for paragraph in doc.paragraphs:
                if "{{zip_code}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{{zip_code}}", qso["zip_code"]
                    )
                if "{{address}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{{address}}", qso["address"]
                    )
                if "{{name}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{name}}", qso["name"])
                if "{{callsign}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{{callsign}}", qso["callsign"]
                    )
                if "{{phone_number}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{{phone_number}}", qso["phone_number"]
                    )
                if "{{image_placeholder}}" in paragraph.text:
                    paragraph.clear()
                    # 插入图片
                    if (
                        DocxGenerator._sign_off_code_generator(qso["sign_off_code"])
                        == 0
                    ):
                        path = (
                            DocxGenerator.path
                            + "/"
                            + DocxGenerator.img_folder
                            + "/"
                            + "qr.png"
                        )
                        paragraph.add_run().add_picture(path, width=Inches(1.0))

            filename = (
                DocxGenerator.path
                + "/"
                + DocxGenerator.mid_folder
                + "/"
                + f"qsl_card_envelope_{idx}.docx"
            )
            doc.save(filename)

        return 0

    @staticmethod
    def _merge():
        paths = DocxGenerator.path + "/" + DocxGenerator.mid_folder
        paths_of_docx = DocxGenerator._find_docx_files_glob(paths)
        base_doc = Document(paths_of_docx[0])
        composer = Composer(base_doc)
        for path in paths_of_docx[1:]:
            doc = Document(path)
            composer.append(doc)

        file_name = datetime.datetime.now().strftime("%Y%m%d_%H-%M-%S")

        output_path = ConfigContext.config["docx_generator"]["output_path"]
        if output_path == "DESKTOP":
            from pathlib import Path

            output_path = str(Path.home()) + "/Desktop"
        file_name = output_path + "/" + file_name + ".docx"
        composer.save(file_name)
        print(f"-{L.get('merge_completed')}{file_name}")

    @staticmethod
    def _find_docx_files_glob(root_folder: str) -> list:
        """
        Only available in python3.10+
        """
        pattern = os.path.join(root_folder, "**", "*.docx")
        return [os.path.abspath(p) for p in glob.glob(pattern, recursive=True)]

    @staticmethod
    def _delete_temp_files(folder_path: str) -> int:
        if not os.path.isdir(folder_path):
            raise FileNotFoundError(f"-{L.get('folder_not_exist')}{folder_path}")
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)  # 删除文件或符号链接
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)  # 递归删除子文件夹
            except Exception as e:
                print(f"-{L.get('delete_failed')}")
                continue

        return True

    @staticmethod
    def generate_envelops_docx(qso: list) -> int:
        DocxGenerator._generator_all_envelopes(qso_list=qso)
        DocxGenerator._merge()
        DocxGenerator._delete_temp_files(
            folder_path=DocxGenerator.path + "/" + DocxGenerator.mid_folder
        )
        return 0

    @staticmethod
    def _sign_off_code_generator(sign_off_code: str) -> int:

        code = (
            ConfigContext.config["global"]["sign_off_url"] + "?token=" + sign_off_code
        )

        qr = qrcode.QRCode(**DocxGenerator.qr_context)
        qr.add_data(code)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        path = DocxGenerator.path + "/" + DocxGenerator.img_folder + "/" + "qr.png"
        img.save(path)
        return 0


if __name__ == "__main__":
    qso_list = [
        {
            "zip_code": "123456",
            "address": "123 Main St",
            "name": "John Doe",
            "callsign": "KD2ABC",
            "sign_off_code": "ABC123",
            "phone_number": "555-555-5555",
        },
        {
            "zip_code": "987654",
            "address": "456 Other St",
            "name": "Jane Smith",
            "callsign": "KD2XYZ",
            "sign_off_code": "XYZ987",
            "phone_number": "555-555-5556",
        },
        {
            "zip_code": "321654",
            "address": "789 Anywhere Ln",
            "name": "Jim Brown",
            "callsign": "KD2DEF",
            "sign_off_code": "DEF321",
            "phone_number": "555-555-5557",
        },
        {
            "zip_code": "654321",
            "address": "321 Nowhere Ave",
            "name": "Jill Jones",
            "callsign": "KD2GHI",
            "sign_off_code": "GHI654",
            "phone_number": "555-555-5558",
        },
        {
            "zip_code": "987321",
            "address": "654 Elsewhere Blvd",
            "name": "Joe Green",
            "callsign": "KD2JKL",
            "sign_off_code": "http://1sdqnwond1odwn1ojwnd1u1j3nf1oui3fn1ou3fn1u3fn1o",
            "phone_number": "555-555-5559",
        },
    ]

    DocxGenerator.generate_envelops_docx(qso_list)
